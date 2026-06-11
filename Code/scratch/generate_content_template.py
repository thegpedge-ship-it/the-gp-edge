import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_heading_1_styled(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    # Left border only (teal color)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="8" w:color="0F766E"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:pBdr>
    ''')
    pPr.append(pBdr)
    
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110) # Teal
    return p

def add_heading_2_styled(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Georgia'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 118, 110) # Teal
    return p

def add_callout_box(doc, title, items, type="teal"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    
    if type == "amber":
        bg_hex = "FFFBEB"
        border_hex = "D97706"
        text_rgb = RGBColor(180, 83, 9)
    else: # teal
        bg_hex = "F0FDFA"
        border_hex = "0F766E"
        text_rgb = RGBColor(15, 118, 110)
        
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border only
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    # Add Title
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(10)
    run_title.font.bold = True
    run_title.font.color.rgb = text_rgb
    
    # Add Items
    for item in items:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.2)
        p_item.paragraph_format.space_before = Pt(2)
        p_item.paragraph_format.space_after = Pt(2)
        
        # Bullet run
        run_bullet = p_item.add_run("•  ")
        run_bullet.font.name = 'Arial'
        run_bullet.font.size = Pt(9.5)
        run_bullet.font.color.rgb = text_rgb
        
        run_text = p_item.add_run(item)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(9.5)
        run_text.font.color.rgb = text_rgb
        
    # Spacer paragraph after table in document
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

def add_styled_table(doc, headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Format Header Row
    hdr_row = table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        set_cell_background(cell, "0F766E") # Teal background
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White text
        
    # Format Data Rows
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC") # Soft grey
            else:
                set_cell_background(cell, "FFFFFF")
                
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # Add bottom borders only
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'''
                <w:tcBorders {nsdecls("w")}>
                    <w:top w:val="none"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>
                    <w:right w:val="none"/>
                </w:tcBorders>
            ''')
            tcPr.append(tcBorders)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(val)
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(71, 85, 105) # Slate
            
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

def setup_page_header_footer(doc, title):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # Configure running header
    header = section.header
    p_hdr = header.paragraphs[0]
    pPr_hdr = p_hdr._p.get_or_add_pPr()
    tabs_hdr = pPr_hdr.find(qn('w:tabs'))
    if tabs_hdr is not None:
        pPr_hdr.remove(tabs_hdr)
    p_hdr.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
    p_hdr.paragraph_format.space_after = Pt(4)
    
    run_hdr_left = p_hdr.add_run(title)
    run_hdr_left.font.name = 'Georgia'
    run_hdr_left.font.size = Pt(8.5)
    run_hdr_left.font.italic = True
    run_hdr_left.font.color.rgb = RGBColor(15, 118, 110) # Teal
    
    run_hdr_tab = p_hdr.add_run("\t")
    
    run_hdr_right = p_hdr.add_run("Medical Condition")
    run_hdr_right.font.name = 'Arial'
    run_hdr_right.font.size = Pt(8.5)
    run_hdr_right.font.color.rgb = RGBColor(148, 163, 184) # Slate Grey
    
    pBdr_hdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="6" w:space="4" w:color="0F766E"/>
        </w:pBdr>
    ''')
    pPr_hdr.append(pBdr_hdr)
    
    # Configure running footer
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    pPr_ftr = p_ftr._p.get_or_add_pPr()
    tabs_ftr = pPr_ftr.find(qn('w:tabs'))
    if tabs_ftr is not None:
        pPr_ftr.remove(tabs_ftr)
    p_ftr.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT)
    p_ftr.paragraph_format.space_before = Pt(4)
    
    run_ftr_left = p_ftr.add_run(f"Synapse Clinical Catalogue — {title}")
    run_ftr_left.font.name = 'Arial'
    run_ftr_left.font.size = Pt(8.5)
    run_ftr_left.font.color.rgb = RGBColor(148, 163, 184) # Slate Grey
    
    pBdr_ftr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/>
        </w:pBdr>
    ''')
    pPr_ftr.append(pBdr_ftr)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_teal_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F0FDFA") # Very light teal background
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Left border only (teal color)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="0F766E"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(15, 118, 110) # Teal
    run.font.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_header(doc, title, subtitle):
    # Setup document geometry
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run(title)
    title_run.font.name = 'Georgia'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 118, 110) # Teal (#0F766E)

    # Document Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate Grey (#64748B)
    
    # Styled bottom border divider
    divider = doc.add_paragraph()
    divider.paragraph_format.space_after = Pt(16)
    pPr = divider._p.get_or_add_pPr()
    pBdr = parse_xml(f'''
        <w:pBdr {nsdecls("w")}>
            <w:bottom w:val="single" w:sz="12" w:space="4" w:color="0F766E"/>
        </w:pBdr>
    ''')
    pPr.append(pBdr)

def build_questions_template(output_path):
    doc = Document()
    add_header(doc, "GP EDGE - Question Import Template", 
               "Use this template to add practice questions to the GP EDGE platform. Fill in the blank fields below.")
    
    add_teal_callout(doc, "GUIDELINES:\n"
                          "1. Modify the blank questions below to include your own practice cases.\n"
                          "2. Keep the tags 'Question X:', 'Options:', 'Correct Answer:', 'High-Yield Rationale:', 'Difficulty:', 'Topic:', 'Subtopic:', 'Tags:' intact.\n"
                          "3. To attach an image, simply paste it inline right under the question text paragraph.")

    for q_num in range(1, 21):
        # Question Block
        q_title = doc.add_paragraph()
        q_title.paragraph_format.space_before = Pt(12)
        q_title.paragraph_format.space_after = Pt(4)
        q_title_run = q_title.add_run(f"Question {q_num}:")
        q_title_run.font.name = 'Georgia'
        q_title_run.font.size = Pt(14)
        q_title_run.font.bold = True
        q_title_run.font.color.rgb = RGBColor(15, 118, 110) # Teal

        q_text = doc.add_paragraph()
        q_text.paragraph_format.space_after = Pt(8)
        q_text_run = q_text.add_run("[Enter question text here]")
        q_text_run.font.name = 'Arial'
        q_text_run.font.size = Pt(11)
        q_text_run.font.color.rgb = RGBColor(30, 41, 59) # Dark Slate

        # Option Header
        opt_header = doc.add_paragraph()
        opt_header.paragraph_format.space_before = Pt(6)
        opt_header.paragraph_format.space_after = Pt(4)
        opt_header_run = opt_header.add_run("Options:")
        opt_header_run.font.name = 'Arial'
        opt_header_run.font.size = Pt(11)
        opt_header_run.font.bold = True
        opt_header_run.font.color.rgb = RGBColor(30, 41, 59)

        # Options
        for letter in ["A", "B", "C", "D"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{letter}) ")
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(71, 85, 105)

        # Correct Answer
        ans_p = doc.add_paragraph()
        ans_p.paragraph_format.space_before = Pt(8)
        ans_p.paragraph_format.space_after = Pt(4)
        run_label = ans_p.add_run("Correct Answer: ")
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(11)
        run_label.font.bold = True
        run_val = ans_p.add_run("")
        run_val.font.name = 'Arial'
        run_val.font.size = Pt(11)
        run_val.font.color.rgb = RGBColor(16, 185, 129) # Emerald Green

        # Rationale Header
        rat_header = doc.add_paragraph()
        rat_header.paragraph_format.space_before = Pt(8)
        rat_header.paragraph_format.space_after = Pt(4)
        rat_header_run = rat_header.add_run("High-Yield Rationale:")
        rat_header_run.font.name = 'Arial'
        rat_header_run.font.size = Pt(11)
        rat_header_run.font.bold = True
        rat_header_run.font.color.rgb = RGBColor(15, 118, 110) # Teal

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        run_desc = p.add_run("[Enter explanation here]")
        run_desc.font.name = 'Arial'
        run_desc.font.size = Pt(10.5)
        run_desc.font.color.rgb = RGBColor(71, 85, 105)

        # Meta parameters
        meta_p = doc.add_paragraph()
        meta_p.paragraph_format.space_before = Pt(8)
        meta_p.paragraph_format.space_after = Pt(24)
        
        meta_items = [
            ("Difficulty", ""),
            ("Topic", ""),
            ("Subtopic", ""),
            ("Tags", "")
        ]
        for i, (k, v) in enumerate(meta_items):
            run_k = meta_p.add_run(f"{k}: ")
            run_k.font.name = 'Arial'
            run_k.font.size = Pt(10)
            run_k.font.bold = True
            run_k.font.color.rgb = RGBColor(100, 116, 139) # Slate
            
            run_v = meta_p.add_run(v)
            run_v.font.name = 'Arial'
            run_v.font.size = Pt(10)
            run_v.font.color.rgb = RGBColor(30, 41, 59)
            
            if i < len(meta_items) - 1:
                meta_p.add_run("\n")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Questions template saved to: {output_path}")

def build_content_template(output_path):
    doc = Document()
    setup_page_header_footer(doc, "[Clinical Guideline/Article Title]")
    
    # Page 1 Header Title block (using the existing add_header style in the body of the first page)
    add_header(doc, "GP EDGE - Medical Content Import Template", 
               "Use this template to draft and import clinical articles and guidelines into the GP EDGE Medical Library.")
    
    add_callout_box(doc, "GUIDELINES:", [
        "Modify the content fields below to write your clinical guideline outline.",
        "Maintain the metadata fields: 'Title:', 'System:', and 'Category:'.",
        "Use numbered Heading 1 sections (such as '1. Section Title') to group your content. They will be styled automatically on import.",
        "To insert styled tables, bullet lists, and callout blocks, format them in Word exactly as demonstrated below."
    ], type="teal")

    # Fields
    fields = [
        ("Title", "[Enter Clinical Article Title Here]"),
        ("System", "[Enter Body System]"),
        ("Category", "[Enter Category]"),
    ]
    
    fields_p = doc.add_paragraph()
    fields_p.paragraph_format.space_after = Pt(12)
    for i, (k, v) in enumerate(fields):
        run_k = fields_p.add_run(f"{k}: ")
        run_k.font.name = 'Arial'
        run_k.font.size = Pt(11)
        run_k.font.bold = True
        run_k.font.color.rgb = RGBColor(15, 118, 110)
        
        run_v = fields_p.add_run(v)
        run_v.font.name = 'Arial'
        run_v.font.size = Pt(11)
        run_v.font.color.rgb = RGBColor(30, 41, 59)
        
        if i < len(fields) - 1:
            fields_p.add_run("\n")

    # 1. Overview
    add_heading_1_styled(doc, "1. Overview")
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(8)
    run_p1 = p1.add_run("[Provide a brief overview of the condition, background context, epidemiology, and general practice relevance here. Keep the description clear, concise, and focused on clinical utility.]")
    run_p1.font.name = 'Arial'
    run_p1.font.size = Pt(10.5)
    run_p1.font.color.rgb = RGBColor(71, 85, 105)

    # 2. Symptoms & Presentation
    add_heading_1_styled(doc, "2. Symptoms")
    for item in [
        "[Key symptom, history question, or clinical feature 1]",
        "[Key symptom, history question, or clinical feature 2]",
        "[Physical examination finding or diagnostic sign 3]",
        "[Physical examination finding or diagnostic sign 4]"
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("•  ").font.color.rgb = RGBColor(71, 85, 105)
        run_t = p.add_run(item)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(71, 85, 105)

    # 3. Diagnosis & Criteria
    add_heading_1_styled(doc, "3. Diagnosis")
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(8)
    run_p3 = p3.add_run("[Specify the diagnostic process, diagnostic tools, investigations, or specific testing criteria below:]")
    run_p3.font.name = 'Arial'
    run_p3.font.size = Pt(10.5)
    run_p3.font.color.rgb = RGBColor(71, 85, 105)

    add_callout_box(doc, "Diagnostic Reference / Criteria Guide", [
        "[Diagnostic criteria step 1, classification rule, or clinical tool referral]",
        "[Diagnostic criteria step 2, classification rule, or clinical tool referral]",
        "[Key laboratory test, radiological investigation, or screening threshold]",
        "[Key laboratory test, radiological investigation, or screening threshold]"
    ], type="teal")

    # 4. Complications
    add_heading_1_styled(doc, "4. Complications")
    add_styled_table(doc, 
        ["Complication / Risk", "Clinical Notes / Prevention"],
        [
            ["[Complication or risk 1]", "[Clinical notes, preventative strategies, or management implications for complication 1]"],
            ["[Complication or risk 2]", "[Clinical notes, preventative strategies, or management implications for complication 2]"],
            ["[Complication or risk 3]", "[Clinical notes, preventative strategies, or management implications for complication 3]"],
            ["[Complication or risk 4]", "[Clinical notes, preventative strategies, or management implications for complication 4]"]
        ]
    )

    # 5. Management & Treatment
    add_heading_1_styled(doc, "5. Management")
    p5 = doc.add_paragraph()
    p5.paragraph_format.space_after = Pt(8)
    run_p5 = p5.add_run("[Detail the care pathway, non-pharmacological adjustments, and drug therapies below. Cite relevant national guidelines where applicable.]")
    run_p5.font.name = 'Arial'
    run_p5.font.size = Pt(10.5)
    run_p5.font.italic = True
    run_p5.font.color.rgb = RGBColor(100, 116, 139)

    add_heading_2_styled(doc, "5a. Non-Pharmacological Management")
    p5a = doc.add_paragraph()
    p5a.paragraph_format.space_after = Pt(8)
    run_p5a = p5a.add_run("[Identify lifestyle modifications, patient education topics, and non-pharmacological strategies to address at each consultation:]")
    run_p5a.font.name = 'Arial'
    run_p5a.font.size = Pt(10.5)
    run_p5a.font.color.rgb = RGBColor(71, 85, 105)

    for item in [
        "[Lifestyle modification, diet/exercise advice, or sleep hygiene parameter 1]",
        "[Lifestyle modification, diet/exercise advice, or sleep hygiene parameter 2]",
        "[Cognitive behavioral therapy, relaxation training, or allied health referral]",
        "[Allied health or physical therapy intervention]"
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("•  ").font.color.rgb = RGBColor(71, 85, 105)
        run_t = p.add_run(item)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(71, 85, 105)

    add_heading_2_styled(doc, "5b. Pharmacological Management")
    p5b1 = doc.add_paragraph()
    p5b1.paragraph_format.space_after = Pt(8)
    run_p5b1 = p5b1.add_run("[Detail pharmacological treatment choices, starting doses, titration schedules, and maximum limits in the table below. Avoid medication overuse risks.]")
    run_p5b1.font.name = 'Arial'
    run_p5b1.font.size = Pt(10.5)
    run_p5b1.font.color.rgb = RGBColor(71, 85, 105)

    add_styled_table(doc,
        ["Drug Class / Example", "Starting Dose", "Maximum Dose", "Titration & Key Side Effects"],
        [
            ["[First-line drug or class 1]", "[Starting dose, route, and frequency]", "[Maximum dose/24h or duration limit]", "[Titration notes, contraindications, and main side effects]"],
            ["[First-line drug or class 2]", "[Starting dose, route, and frequency]", "[Maximum dose/24h or duration limit]", "[Titration notes, contraindications, and main side effects]"],
            ["[Second-line drug or class 3]", "[Starting dose, route, and frequency]", "[Maximum dose/24h or duration limit]", "[Titration notes, contraindications, and main side effects]"],
            ["[Alternative / Adjoint class 4]", "[Starting dose, route, and frequency]", "[Maximum dose/24h or duration limit]", "[Titration notes, contraindications, and main side effects]"]
        ]
    )

    # 6. Warning Signs & Red Flags
    add_heading_1_styled(doc, "6. Warning Signs")
    p6_desc = doc.add_paragraph()
    p6_desc.paragraph_format.space_after = Pt(6)
    run_p6_desc = p6_desc.add_run("[Identify clinical red flags, alarm symptoms, or indicators that require immediate escalation or referral:]")
    run_p6_desc.font.name = 'Arial'
    run_p6_desc.font.size = Pt(10.5)
    run_p6_desc.font.color.rgb = RGBColor(71, 85, 105)

    add_callout_box(doc, "Warning Signs & Emergency Red Flags", [
        "[Warning sign or alarm symptom requiring urgent secondary workup 1]",
        "[Warning sign or alarm symptom requiring urgent secondary workup 2]",
        "[Contraindication or clinical criteria requiring immediate specialist escalation]"
    ], type="amber")

    # 7. When to Refer
    add_heading_1_styled(doc, "7. When to Refer")
    p7_desc = doc.add_paragraph()
    p7_desc.paragraph_format.space_after = Pt(6)
    run_p7_desc = p7_desc.add_run("[List specific criteria for outpatient specialist referral, allied health coordination, or diagnostic support:]")
    run_p7_desc.font.name = 'Arial'
    run_p7_desc.font.size = Pt(10.5)
    run_p7_desc.font.color.rgb = RGBColor(71, 85, 105)

    for item in [
        "[Referral criteria 1]",
        "[Referral criteria 2]",
        "[Referral criteria 3]"
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("•  ").font.color.rgb = RGBColor(71, 85, 105)
        run_t = p.add_run(item)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(71, 85, 105)

    # 8. Prognosis
    add_heading_1_styled(doc, "8. Prognosis")
    for item in [
        "[Expected clinical course and recovery timeline for acute presentations]",
        "[Long-term outcomes, recurrence rates, and risk of progression to chronic disease]",
        "[Expected improvement with lifestyle modifications and secondary prevention compliance]"
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("•  ").font.color.rgb = RGBColor(71, 85, 105)
        run_t = p.add_run(item)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(71, 85, 105)

    # 9. Resources
    add_heading_1_styled(doc, "9. Resources")
    p9 = doc.add_paragraph()
    p9.paragraph_format.space_after = Pt(4)
    run_p9 = p9.add_run("For Health Professionals")
    run_p9.font.name = 'Arial'
    run_p9.font.size = Pt(11)
    run_p9.font.bold = True
    run_p9.font.color.rgb = RGBColor(15, 118, 110)

    for item in [
        "[Primary national guidelines or reference handbook 1]",
        "[Secondary reference website or decision support tool link 2]"
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("•  ").font.color.rgb = RGBColor(71, 85, 105)
        run_t = p.add_run(item)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(71, 85, 105)

    p9_2 = doc.add_paragraph()
    p9_2.paragraph_format.space_before = Pt(8)
    p9_2.paragraph_format.space_after = Pt(4)
    run_p9_2 = p9_2.add_run("For Patients")
    run_p9_2.font.name = 'Arial'
    run_p9_2.font.size = Pt(11)
    run_p9_2.font.bold = True
    run_p9_2.font.color.rgb = RGBColor(15, 118, 110)

    for item in [
        "[Patient information factsheet, support organization website, or mobile application 1]",
        "[Patient information factsheet, support organization website, or mobile application 2]"
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(3)
        p.add_run("•  ").font.color.rgb = RGBColor(71, 85, 105)
        run_t = p.add_run(item)
        run_t.font.name = 'Arial'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = RGBColor(71, 85, 105)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Content template saved to: {output_path}")


def build_autofill_template(output_path):
    doc = Document()
    add_header(doc, "GP EDGE - Autofill Import Template", 
               "Use this template to create custom SOAP note autofill templates.")
    
    add_teal_callout(doc, "GUIDELINES:\n"
                          "1. A shortcut is a text trigger starting with a semicolon (e.g. ;shortcut).\n"
                          "2. Set the 'System' to group the template under a specific system tab.\n"
                          "3. Include text placeholder fields inside brackets like [Placeholder] so they can be easily tabbed through in the consultation editor.")

    # Fields
    fields = [
        ("Template Name", "[Enter Template Name]"),
        ("Shortcut", "[Enter Shortcut]"),
        ("System", "[Enter Body System]"),
    ]
    
    fields_p = doc.add_paragraph()
    fields_p.paragraph_format.space_after = Pt(12)
    for i, (k, v) in enumerate(fields):
        run_k = fields_p.add_run(f"{k}: ")
        run_k.font.name = 'Arial'
        run_k.font.size = Pt(11)
        run_k.font.bold = True
        run_k.font.color.rgb = RGBColor(15, 118, 110)
        
        run_v = fields_p.add_run(v)
        run_v.font.name = 'Arial'
        run_v.font.size = Pt(11)
        run_v.font.color.rgb = RGBColor(30, 41, 59)
        
        if i < len(fields) - 1:
            fields_p.add_run("\n")

    # SOAP notes
    for section_name, text in [
        ("Subjective", "Patient presents for [Reason for Visit]. Describes [Symptom description] lasting for [Duration]. Symptoms are [Severity]. Worsened by [Triggers] and relieved by [Relieving factors]. Denies [Relevant red flag symptoms]."),
        ("Objective", "Vitals: BP [Blood Pressure] mmHg, HR [Heart Rate] bpm, Temp [Temperature] °C.\n"
                      "Examination: Abdomen [Findings]. Chest [Findings]. Cardiovascular [Findings]. No signs of [Alarm symptoms]."),
        ("Assessment", "[Diagnosis/Impression]. Controlled/uncontrolled, with [Complications/Comorbidities]. No red flags present at this time."),
        ("Plan", "1. Lifestyle modification: [Recommendations].\n"
                 "2. Start Pharmacotherapy: [Medication details].\n"
                 "3. Order investigations: [Investigations].\n"
                 "4. Review in [Review time]. Advised to return immediately if [Warning signs] develops.")
    ]:
        sec = doc.add_paragraph()
        sec.paragraph_format.space_before = Pt(8)
        sec.paragraph_format.space_after = Pt(4)
        run_sec = sec.add_run(section_name)
        run_sec.font.name = 'Georgia'
        run_sec.font.size = Pt(14)
        run_sec.font.bold = True
        run_sec.font.color.rgb = RGBColor(15, 118, 110)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run_p = p.add_run(text)
        run_p.font.name = 'Arial'
        run_p.font.size = Pt(10.5)
        run_p.font.color.rgb = RGBColor(71, 85, 105)

    doc.save(output_path)
    print(f"Autofill template saved to: {output_path}")

if __name__ == "__main__":
    base_dir = "d:\\Projects\\GP EDGE\\Code\\public\\templates"
    build_questions_template(os.path.join(base_dir, "question_template.docx"))
    build_content_template(os.path.join(base_dir, "medical_content_template.docx"))
    build_autofill_template(os.path.join(base_dir, "autofill_template.docx"))
